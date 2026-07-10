import base64
import io
import re
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from github import Github
from odf import teletype
from odf.opendocument import load
from odf.text import P
from pypdf import PdfReader


st.set_page_config(page_title="Netflix Data Editor", layout="wide")
st.title("📺 Watch History Lookup & Editor")

SUPPORTED_TYPES = ["csv", "txt", "odt", "docx", "pdf", "xlsx", "xls", "json"]
DATE_AT_END = re.compile(
    r"^(?P<Title>.*?)[\s,;-]+(?P<Date>\d{1,4}[/-]\d{1,2}[/-]\d{1,4})\s*$"
)


def text_to_dataframe(text):
    """Read table-like text, or turn a plain watched-program list into rows."""
    text = text.replace("\ufeff", "").strip()
    if not text:
        raise ValueError("The document does not contain readable text.")

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    # Netflix exports and hand-built lists are commonly comma/tab separated.
    for separator in (",", "\t", ";", "|"):
        try:
            candidate = pd.read_csv(
                io.StringIO(text), sep=separator, engine="python", on_bad_lines="skip"
            )
            if len(candidate.columns) > 1:
                return candidate
        except (pd.errors.ParserError, UnicodeDecodeError):
            pass

    # Plain documents often contain one title per line, optionally ending in a date.
    rows = []
    for line in lines:
        match = DATE_AT_END.match(line)
        rows.append(match.groupdict() if match else {"Title": line, "Date": None})
    return pd.DataFrame(rows)


def extract_odt(content):
    document = load(io.BytesIO(content))
    return "\n".join(
        teletype.extractText(paragraph)
        for paragraph in document.getElementsByType(P)
    )


def extract_docx(content):
    document = Document(io.BytesIO(content))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        lines.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(lines)


def extract_pdf(content):
    reader = PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_document(content, filename):
    extension = Path(filename).suffix.lower()

    if extension == ".csv":
        return pd.read_csv(io.BytesIO(content), on_bad_lines="skip")
    if extension == ".txt":
        return text_to_dataframe(content.decode("utf-8-sig", errors="replace"))
    if extension == ".odt":
        return text_to_dataframe(extract_odt(content))
    if extension == ".docx":
        return text_to_dataframe(extract_docx(content))
    if extension == ".pdf":
        return text_to_dataframe(extract_pdf(content))
    if extension in (".xlsx", ".xls"):
        return pd.read_excel(io.BytesIO(content))
    if extension == ".json":
        try:
            return pd.read_json(io.BytesIO(content))
        except ValueError:
            return pd.json_normalize(pd.read_json(io.BytesIO(content), typ="series"))

    raise ValueError(
        f"Unsupported file type '{extension or 'unknown'}'. "
        f"Choose one of: {', '.join(SUPPORTED_TYPES)}."
    )


def clean_dataframe(df):
    df = df.copy()
    df.columns = (
        df.columns.astype(str)
        .str.strip()
        .str.replace('"', "", regex=False)
        .str.replace("'", "", regex=False)
        .str.replace("ï»¿", "", regex=False)
    )
    if "Date" in df.columns:
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    return df


@st.cache_data
def load_github_data():
    github = Github(st.secrets["GITHUB_TOKEN"])
    repo = github.get_repo(st.secrets["REPO_NAME"])
    github_file = repo.get_contents(st.secrets["FILE_PATH"])
    content = base64.b64decode(github_file.content)
    return clean_dataframe(read_document(content, st.secrets["FILE_PATH"])), github_file.sha


source = st.radio("Document source", ["Upload a document", "GitHub file"], horizontal=True)
df = None
sha = None

if source == "Upload a document":
    uploaded_file = st.file_uploader(
        "Choose a watch-history document", type=SUPPORTED_TYPES
    )
    if uploaded_file is not None:
        try:
            df = clean_dataframe(read_document(uploaded_file.getvalue(), uploaded_file.name))
        except Exception as error:
            st.error(f"Could not read {uploaded_file.name}: {error}")
else:
    try:
        df, sha = load_github_data()
    except Exception as error:
        st.error(f"Error loading data from GitHub: {error}")

if df is not None:
    st.subheader("Search Watch History")

    title_column = next(
        (
            column
            for column in df.columns
            if column.lower() in ("title", "program", "show", "name")
        ),
        df.columns[0] if len(df.columns) else None,
    )
    search_title = st.text_input(
        "Search by title",
        placeholder="Enter a movie, show, or episode title",
    ).strip()

    if search_title and title_column is not None:
        matches = df[
            df[title_column]
            .fillna("")
            .astype(str)
            .str.contains(search_title, case=False, regex=False)
        ].copy()

        if matches.empty:
            st.warning(f'No watched programs found matching "{search_title}".')
        else:
            date_column = next(
                (column for column in matches.columns if column.lower() == "date"),
                None,
            )
            if date_column is not None:
                matches[date_column] = pd.to_datetime(
                    matches[date_column], errors="coerce"
                ).dt.strftime("%B %d, %Y").str.replace(" 0", " ", regex=False)
                matches[date_column] = matches[date_column].fillna("Date unavailable")

            st.success(
                f"Found {len(matches)} watched record"
                f"{'s' if len(matches) != 1 else ''}."
            )
            result_columns = [title_column]
            if date_column is not None and date_column != title_column:
                result_columns.append(date_column)
            st.dataframe(matches[result_columns], use_container_width=True, hide_index=True)

    st.subheader("Edit Watch History")
    edited_df = st.data_editor(df, num_rows="dynamic", use_container_width=True)
    csv_bytes = edited_df.to_csv(index=False).encode("utf-8")

    if source == "Upload a document":
        st.download_button(
            "Download Edited CSV",
            data=csv_bytes,
            file_name="watch_history_edited.csv",
            mime="text/csv",
        )
    elif st.button("Save Changes to GitHub"):
        try:
            github = Github(st.secrets["GITHUB_TOKEN"])
            repo = github.get_repo(st.secrets["REPO_NAME"])
            repo.update_file(
                path=st.secrets["FILE_PATH"],
                message="Updated watch history via Streamlit",
                content=csv_bytes,
                sha=sha,
            )
            load_github_data.clear()
            st.success("Changes saved to GitHub!")
        except Exception as error:
            st.error(f"Error saving file: {error}")
elif source == "Upload a document":
    st.info("Upload a document to view and edit its watched programs.")
else:
    st.info("No data loaded. Check the GitHub settings in secrets.toml.")
